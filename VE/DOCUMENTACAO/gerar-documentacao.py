# -*- coding: utf-8 -*-
"""Gera VE-Projeto-Documentacao.docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

OUT = r"c:\Users\z005027j\Desktop\VE\DOCUMENTACAO\VE-Projeto-Documentacao.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def h3(t):
    return doc.add_heading(t, level=3)

def para(t, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    if bold:
        r.bold = True
    return p

def bullet(items):
    for i in items:
        doc.add_paragraph(i, style="List Bullet")

def numbered(items):
    for i in items:
        doc.add_paragraph(i, style="List Number")

# Capa
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Projeto VE\nBriefing candeias.dev + Cliente Vénusia")
r.bold = True
r.font.size = Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"Documentação completa · {date.today().strftime('%d/%m/%Y')}").italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Autor: candeias.dev web-solutions\nWhatsApp receção: 938 024 721\nPasta: C:\\Users\\z005027j\\Desktop\\VE")

doc.add_page_break()

# 1
h1("1. Resumo do projeto")
para(
    "Este projeto foi criado para a cliente Vénusia (referência visual: venusia.monfily.com) "
    "e, em paralelo, um sistema de briefing reutilizável para todos os clientes da candeias.dev. "
    "O objetivo principal é recolher informação estruturada para desenvolvimento de websites "
    "através de um formulário online com 35 perguntas, envio direto por WhatsApp, "
    "disponível em português e inglês."
)

# 2
h1("2. O que foi construído")
h2("2.1 Formulário de briefing (versão oficial)")
bullet([
    "35 perguntas organizadas por secções + secção personalizada + notas",
    "Tema visual preto estilo Cursor (fundo #000, grelha, cards escuros, fonte Inter)",
    "Cabeçalho: Briefing · candeias.dev · web-solutions (link Instagram)",
    "Alternância de idioma PT | EN no topo (i18n em JavaScript)",
    "Perguntas com hints e exemplos; cliente pode responder «não sei» ou «a definir»",
    "Envio WhatsApp para 351938024721 com fallback copiar texto + abrir chat",
    "Marcador de versão visível: «briefing v2»",
    "Ficheiro principal: deploy\\index.html",
])

h2("2.2 Publicação online (Netlify)")
bullet([
    "Plataforma: Netlify (plano gratuito)",
    "URL final: https://candeias-briefing.netlify.app/",
    "Inglês: https://candeias-briefing.netlify.app/?lang=en ou /en.html",
    "Nome do projeto no Netlify: candeias-briefing",
    "Deploy inicial via Netlify Drop; atualizações via arrastar pasta deploy",
    "Ficheiro _headers para evitar cache antigo do index.html",
])
h2("2.3 Pacotes offline (legado)")
bullet([
    "PACOTES\\Briefing-Website-PT.zip e EN.zip — Windows, extrair + .bat",
    "PACOTES\\Briefing-Website-PT-IPHONE.html — iPhone sem ZIP",
    "Versão bege antiga; preferir sempre o link Netlify",
])
h2("2.4 Mensagens WhatsApp prontas")
bullet([
    "Pasta MENSAGENS\\ com textos para primeiro contacto, lembrete, confirmação",
    "05b-whatsapp-colar-direto.txt — mensagem principal com link clicável",
    "Guia LINK-CLICAVEL-WHATSAPP.txt — regras para link ficar azul",
])

# 3
h1("3. Cronologia do que aconteceu")
numbered([
    "Utilizador reiniciou o projeto do zero com estrutura organizada.",
    "Criação dos formulários PT/EN iniciais (versão bege) em BRIEFING e PACOTES.",
    "Evolução para versão preta unificada em deploy\\ com 35 perguntas e marca candeias.dev.",
    "Problema de preview local: ERR_EMPTY_RESPONSE em localhost:5500.",
    "Causa: Python http.server em IPv6 (::); browser em IPv4.",
    "Solução: preview.bat com --bind 127.0.0.1 e URL http://127.0.0.1:5500.",
    "Publicação no Netlify (primeiro link: exquisite-cajeta-30c006.netlify.app).",
    "Upload acidental da versão bege antiga em vez de deploy\\.",
    "Correção: redeploy da pasta deploy + marcador briefing v2.",
    "Renomeação do site para candeias-briefing.netlify.app.",
    "Ajuste de mensagens WhatsApp para link clicável (URL em linha isolada, sem emoji).",
    "Organização final da pasta + este documento Word.",
])

# 4
h1("4. Estrutura da pasta (após organização)")
para("C:\\Users\\z005027j\\Desktop\\VE\\", bold=True)
structure = """
VE\\
├── LEIA-ME.txt                 → Índice rápido do projeto
├── preview.bat                 → Servidor local porta 5500
├── atualizar-netlify.bat       → Abre deploy + Netlify para redeploy
├── abrir-pasta-deploy.bat      → Abre pasta deploy
├── package.json                → npm run preview (alternativa)
│
├── deploy\\                     ★ PUBLICAR NO NETLIFY (versão preta)
│   ├── index.html
│   ├── en.html
│   ├── _headers
│   ├── COMO-PUBLICAR-LINK.txt
│   └── PREVIEW-LOCAL.txt
│
├── MENSAGENS\\                  Textos WhatsApp
│   ├── 00-INDICE.txt
│   ├── 05b-whatsapp-colar-direto.txt  (principal)
│   └── …
│
├── PACOTES\\                    ZIP + iPhone (offline, bege)
│   └── LEIA-ME.txt
│
├── DOCUMENTACAO\\
│   ├── VE-Projeto-Documentacao.docx  (este ficheiro)
│   └── gerar-documentacao.py
│
└── _ARQUIVO\\
    └── briefing-versao-bege\\   Versão antiga (não publicar)
"""
para(structure)

# 5
h1("5. Como usar no dia a dia")
h2("5.1 Enviar briefing a um cliente")
numbered([
    "Abrir MENSAGENS\\05b-whatsapp-colar-direto.txt",
    "Substituir saudação pelo nome do cliente",
    "Copiar tudo (Ctrl+A, Ctrl+C) e colar no WhatsApp",
    "Enviar — o link https://candeias-briefing.netlify.app/ deve ficar azul",
    "Cliente preenche e toca em «Enviar no WhatsApp»",
    "Recebe no 938 024 721",
])
h2("5.2 Testar localmente")
numbered([
    "Duplo clique em preview.bat",
    "Abrir http://127.0.0.1:5500/",
    "Confirmar fundo preto e «briefing v2»",
    "Parar com Ctrl+C",
])
h2("5.3 Atualizar o Netlify")
numbered([
    "Editar ficheiros em deploy\\",
    "Duplo clique em atualizar-netlify.bat",
    "No Netlify → Deploys → arrastar pasta deploy inteira",
    "Testar no telemóvel (fechar separador antigo se necessário)",
])

# 6
h1("6. Primeira mensagem para a Vénusia (modelo)")
msg = """Olá Vénusia!

Obrigada pelo contacto. Para avançarmos com o website de forma organizada, preparei um briefing online — demora cerca de 15 a 20 minutos.

https://candeias-briefing.netlify.app/

Como funciona:
1) Abra o link no telemóvel ou computador
2) Responda às perguntas (onde não souber, pode escrever «não sei» ou «a definir»)
3) No final, toque em Enviar no WhatsApp — recebo tudo automaticamente

Fotos, logótipo e outros ficheiros: envie à parte por WhatsApp quando tiver.

Qualquer dúvida, estou por aqui.

Obrigada!"""
p = doc.add_paragraph()
p.add_run(msg).font.size = Pt(10)

# 7
h1("7. Problemas resolvidos")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Problema"
hdr[1].text = "Causa"
hdr[2].text = "Solução"
rows = [
    ("Preview localhost não abre", "Servidor só IPv6", "preview.bat com --bind 127.0.0.1"),
    ("Netlify mostra versão bege", "Upload pasta errada", "Só arrastar deploy\\"),
    ("Link WhatsApp não fica azul", "Emoji na mesma linha do URL", "URL sozinho numa linha"),
    ("Botão WhatsApp tapado no mobile", "Barra fixa sobreposta", "CSS z-index e padding"),
    ("iPhone não abre HTML no ZIP", "Limitação iOS", "Link Netlify ou .html solto"),
    ("Reenvio WhatsApp falha", "Limite URL / pop-up", "Copiar texto + abrir chat"),
]
for prob, causa, sol in rows:
    row = table.add_row().cells
    row[0].text = prob
    row[1].text = causa
    row[2].text = sol

# 8
h1("8. Dados técnicos")
bullet([
    "WhatsApp: WHATSAPP_NUMBER = 351938024721",
    "Formulário: HTML + CSS + JavaScript (sem framework)",
    "Preview: Python 3 http.server porta 5500",
    "Alternativa: npm run preview (npx serve)",
    "Netlify: hospedagem estática gratuita",
    "Idiomas: objeto T em JS; ?lang=en na URL",
])

# 9
h1("9. Próximos passos sugeridos")
bullet([
    "Aguardar briefing preenchido pela Vénusia",
    "Desenvolver site Vénusia com material recebido",
    "Opcional: aplicar mesmo tema preto Cursor ao site final",
    "Sincronizar PACOTES\\ iPhone com deploy\\ se necessário",
    "Dominio personalizado no Netlify (opcional, pago)",
])

# 10
h1("10. Contactos e links")
bullet([
    "Link briefing: https://candeias-briefing.netlify.app/",
    "Instagram: https://www.instagram.com/candeias.dev",
    "Netlify Drop: https://app.netlify.com/drop",
    "Netlify painel: https://app.netlify.com/",
])

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("— Fim do documento —").italic = True

doc.save(OUT)
print("OK:", OUT)
